#!/usr/bin/env python3
"""
Polished family-summary builder (S-cleanup). Reads a source .md, folds Sanskrit to
plain READABLE English (family-friendly: Panini, shunya, ashta-dik -- NOT diacritics,
NOT Harvard-Kyoto), and emits a cleanly-formatted .docx (title block, coloured
headings, banded tables w/ repeating header, 1.15 spacing, proper margins).
Then convert the .docx -> .pdf via Word (docx2pdf) separately.
"""
import os
import re
import unicodedata
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SB = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "Science_Background")  # 01_CHIP_A0S/Science_Background (reorg Phase 2)
SRC = os.path.join(SB, "archive", "Ansh_108_Core_For_Family_Plain_Summary.md")  # archived content base
OUT_MD = os.path.join(SB, "Ansh_108_Core_Family_Summary.md")
OUT_DOCX = os.path.join(SB, "Ansh_108_Core_Family_Summary.docx")

ACCENT = RGBColor(0x1F, 0x49, 0x6D)   # deep blue for the title
H1COL = RGBColor(0x2E, 0x74, 0xB5)    # blue for section heads
H2COL = RGBColor(0x40, 0x40, 0x40)    # dark grey for sub-heads
GREY = RGBColor(0x60, 0x60, 0x60)

# Sanskrit/IAST + Devanagari -> plain readable English (family-friendly)
READ = {
    "ā": "a", "ī": "i", "ū": "u", "ṛ": "ri", "ṝ": "ri", "ḷ": "l", "ḹ": "l",
    "ṅ": "n", "ñ": "n", "ṭ": "t", "ḍ": "d", "ṇ": "n", "ś": "sh", "ṣ": "sh",
    "ṃ": "m", "ṁ": "m", "ḥ": "h",
    "Ā": "A", "Ī": "I", "Ū": "U", "Ṛ": "Ri", "Ṅ": "N", "Ñ": "N", "Ṭ": "T",
    "Ḍ": "D", "Ṇ": "N", "Ś": "Sh", "Ṣ": "Sh", "Ṃ": "M", "Ḥ": "H",
    "ï": "i", "î": "i", "é": "e", "è": "e", "ê": "e", "ä": "a", "â": "a",
    "ö": "o", "ô": "o", "ü": "u", "û": "u", "ç": "c",
    "अंश": "amsha",
    # avoid heavy emoji-font embedding (Segoe UI Emoji ~MB): use plain glyphs
    "✅": "✓", "✔": "✓", "❌": "x", "✗": "x",
}
INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def fold(t):
    t = unicodedata.normalize("NFC", t)
    for k, v in READ.items():
        t = t.replace(k, v)
    return t


def add_runs(par, text):
    text = fold(text)
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        seg = m.group(0)
        if seg.startswith("**"):
            par.add_run(seg[2:-2]).bold = True
        else:
            par.add_run(seg[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:tblHeader")
    h.set(qn("w:val"), "true")
    trPr.append(h)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main():
    with open(SRC, "r", encoding="utf-8") as fh:
        raw = fh.read()
    # write the readable-English source alongside (the new .md)
    with open(OUT_MD, "w", encoding="utf-8") as fh:
        fh.write(fold(raw))

    lines = raw.split("\n")
    doc = Document()
    sec = doc.sections[0]
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.9))
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    n_h3 = 0
    header_block = True       # everything before the first '## ' is the title block
    i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()
        if not s:
            i += 1
            continue

        # ---- tables ----
        if s.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = [split_row(b) for b in block]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=0, cols=ncols)
            for style in ("Medium Shading 1 Accent 1", "Light Grid Accent 1", "Table Grid"):
                try:
                    tbl.style = style
                    break
                except Exception:
                    continue
            for ri, r in enumerate(rows):
                cells = tbl.add_row().cells
                for ci in range(ncols):
                    txt = r[ci] if ci < len(r) else ""
                    p = cells[ci].paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    if ri == 0:
                        run = p.add_run(fold(re.sub(r"\*\*", "", txt)))
                        run.bold = True
                    else:
                        add_runs(p, txt)
            repeat_header(tbl.rows[0])
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        if s.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(fold(s[2:]))
            r.bold = True
            r.font.size = Pt(23)
            r.font.color.rgb = ACCENT
        elif s.startswith("## "):
            header_block = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(fold(s[3:]))
            r.bold = True
            r.font.size = Pt(14.5)
            r.font.color.rgb = H1COL
        elif s.startswith("### "):
            n_h3 += 1
            p = doc.add_paragraph()
            if n_h3 == 1:            # the subtitle, right under the title
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(fold(s[4:]))
                r.italic = True
                r.font.size = Pt(12.5)
                r.font.color.rgb = GREY
            else:
                p.paragraph_format.space_before = Pt(8)
                r = p.add_run(fold(s[4:]))
                r.bold = True
                r.font.size = Pt(12)
                r.font.color.rgb = H2COL
        elif s == "---":
            doc.add_paragraph().paragraph_format.space_after = Pt(2)   # quiet spacer
        elif s.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            run_par = p
            add_runs(run_par, s[2:])
            for rr in run_par.runs:
                rr.italic = True
                rr.font.color.rgb = GREY
        elif s.startswith("- "):
            add_runs(doc.add_paragraph(style="List Bullet"), s[2:])
        else:
            p = doc.add_paragraph()
            if header_block:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(p, s)
        i += 1

    doc.save(OUT_DOCX)
    print("WROTE", os.path.abspath(OUT_MD))
    print("WROTE", os.path.abspath(OUT_DOCX))


if __name__ == "__main__":
    main()
