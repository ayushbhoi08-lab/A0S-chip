#!/usr/bin/env python3
"""
Build the family-summary .docx from the plain-words .md (single source of truth).
Handles the subset of Markdown the summary uses: #/##/### headings, --- rules,
> blockquotes, - bullets, | pipe tables |, paragraphs, and inline **bold**/*italic*.

  python build_family_doc.py                 -> rich .docx (keeps Sanskrit/diacritics)
  python build_family_doc.py --ascii --out X -> ASCII-folded .docx (lean, glyph-free PDF)

Run from the Artifacts dir; default output is the .docx next to the .md.
"""
import os
import re
import sys
import unicodedata
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "..", "Ansh_108_Core_For_Family_Plain_Summary.md")
DEFAULT_OUT = os.path.join(HERE, "..", "Ansh_108_Core_For_Family_Plain_Summary.docx")

ASCII = "--ascii" in sys.argv
OUT = DEFAULT_OUT
if "--out" in sys.argv:
    OUT = sys.argv[sys.argv.index("--out") + 1]

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")

# explicit replacements for symbols / Devanagari that NFKD can't fold to ASCII
_REPL = {
    "अंश": "amsha", "✅": "[OK]", "→": "->", "×": "x", "·": " - ",
    "—": "-", "–": "-", "―": "-", "≈": "~", "…": "...", "•": "-",
    "“": '"', "”": '"', "‘": "'", "’": "'", "°": " deg", " ": " ",
}


def fold(t):
    """Fold to plain ASCII when --ascii: transliterate diacritics, drop Devanagari."""
    if not ASCII:
        return t
    for k, v in _REPL.items():
        t = t.replace(k, v)
    t = unicodedata.normalize("NFKD", t)
    return t.encode("ascii", "ignore").decode("ascii")


def add_runs(par, text):
    text = fold(text)
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        seg = m.group(0)
        if seg.startswith("**"):
            par.add_run(seg[2:-2]).bold = True
        else:
            par.add_run(seg[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main():
    with open(SRC, "r", encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()
        if not s:
            i += 1
            continue

        if s.startswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            rows = [split_row(b) for b in block]
            rows = [r for r in rows if not all(set(c) <= set("-: ") for c in r)]
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=0, cols=ncols)
            try:
                tbl.style = "Light Grid Accent 1"
            except Exception:
                tbl.style = "Table Grid"
            for ri, r in enumerate(rows):
                cells = tbl.add_row().cells
                for ci in range(ncols):
                    txt = r[ci] if ci < len(r) else ""
                    p = cells[ci].paragraphs[0]
                    if ri == 0:
                        p.add_run(fold(re.sub(r"\*\*", "", txt))).bold = True
                    else:
                        add_runs(p, txt)
            doc.add_paragraph()
            continue

        if s.startswith("# "):
            doc.add_heading(fold(s[2:]), level=0)
        elif s.startswith("## "):
            doc.add_heading(fold(s[3:]), level=1)
        elif s.startswith("### "):
            doc.add_heading(fold(s[4:]), level=2)
        elif s == "---":
            doc.add_paragraph("-" * 24).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif s.startswith("> "):
            p = doc.add_paragraph()
            try:
                p.style = doc.styles["Intense Quote"]
            except Exception:
                p.paragraph_format.left_indent = Pt(18)
            add_runs(p, s[2:])
        elif s.startswith("- "):
            add_runs(doc.add_paragraph(style="List Bullet"), s[2:])
        else:
            add_runs(doc.add_paragraph(), s)
        i += 1

    doc.save(OUT)
    print(("ASCII " if ASCII else "RICH  ") + "WROTE " + os.path.abspath(OUT))


if __name__ == "__main__":
    main()
